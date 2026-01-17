"""
Word Document Exporter - Export processed documents to .docx format.

Creates Word documents with text, tables, images, and SVG (as images).
"""

import io
import re
import tempfile
from pathlib import Path
from typing import Optional, List, Dict, Any
import logging

logger = logging.getLogger(__name__)


class WordExporter:
    """
    Export processed document content to Word (.docx) format.

    Creates a replica of the original document with:
    - OCR text content
    - Tables (from extracted table data)
    - Images (original or regenerated)
    - Charts (SVG converted to images)
    """

    def __init__(self):
        """Initialize the Word exporter."""
        self._document = None

    def export(
        self,
        document_dir: str | Path,
        output_path: str | Path,
        structured_data: Optional[Dict[str, Any]] = None,
        include_regenerated_images: bool = True,
        include_original_images: bool = False
    ) -> str:
        """
        Export document content to Word format.

        Args:
            document_dir: Directory containing processed document files
            output_path: Path for output .docx file
            structured_data: Optional pre-loaded structured document data
            include_regenerated_images: Include AI-regenerated images
            include_original_images: Include original extracted images

        Returns:
            Path to the created Word document
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        document_dir = Path(document_dir)
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Load structured data if not provided
        if structured_data is None:
            structured_data = self._load_structured_data(document_dir)

        # Create Word document
        doc = Document()

        # Add title
        title = structured_data.get("metadata", {}).get("title") or document_dir.name
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        self._add_metadata_section(doc, structured_data)

        # Process content page by page
        pages_content = structured_data.get("content", {}).get("pages", [])
        tables_data = structured_data.get("extracted_data", {}).get("tables", [])
        charts_data = structured_data.get("extracted_data", {}).get("charts", [])
        images_data = structured_data.get("extracted_data", {}).get("images", [])

        # Group content by page
        for page in pages_content:
            page_num = page.get("page_number", 0)

            # Add page separator
            if page_num > 1:
                doc.add_page_break()

            # Add page heading
            doc.add_heading(f"Page {page_num}", level=1)

            # Add page text
            text = page.get("text", "")
            if text:
                self._add_text_content(doc, text)

            # Add tables for this page
            page_tables = [t for t in tables_data if t.get("page_number") == page_num]
            for table in page_tables:
                self._add_table(doc, table)

            # Add charts for this page (SVG as images)
            page_charts = [c for c in charts_data if c.get("page_number") == page_num]
            for chart in page_charts:
                self._add_chart(doc, chart, document_dir)

            # Add images for this page
            if include_regenerated_images:
                self._add_regenerated_images(doc, page_num, document_dir)

            if include_original_images:
                page_images = [i for i in images_data if i.get("page_number") == page_num]
                for image in page_images:
                    self._add_image(doc, image, document_dir)

        # Save document
        doc.save(str(output_path))
        logger.info(f"Word document saved: {output_path}")

        return str(output_path)

    def _load_structured_data(self, document_dir: Path) -> Dict[str, Any]:
        """Load structured document data from JSON file."""
        import json

        json_path = document_dir / "document.json"
        if json_path.exists():
            with open(json_path, "r", encoding="utf-8") as f:
                return json.load(f)

        # Fallback: build basic structure from available files
        return self._build_basic_structure(document_dir)

    def _build_basic_structure(self, document_dir: Path) -> Dict[str, Any]:
        """Build basic structure from available files."""
        import json

        structure = {
            "metadata": {"title": document_dir.name},
            "content": {"pages": []},
            "extracted_data": {"tables": [], "charts": [], "images": []}
        }

        # Load text files
        txt_dir = document_dir / "txt"
        if txt_dir.exists():
            for txt_file in sorted(txt_dir.glob("page_*.txt")):
                match = re.search(r'page_(\d+)', txt_file.stem)
                page_num = int(match.group(1)) if match else len(structure["content"]["pages"]) + 1

                with open(txt_file, "r", encoding="utf-8") as f:
                    text = f.read()

                structure["content"]["pages"].append({
                    "page_number": page_num,
                    "text": text
                })

        # Load tables
        json_dir = document_dir / "json"
        if json_dir.exists():
            for json_file in sorted(json_dir.glob("*_table_*.json")):
                with open(json_file, "r", encoding="utf-8") as f:
                    table_data = json.load(f)
                structure["extracted_data"]["tables"].append(table_data)

        # Load charts
        svg_dir = document_dir / "svg"
        if svg_dir.exists():
            for svg_file in sorted(svg_dir.glob("*_chart_*.svg")):
                match = re.search(r'page_(\d+)_chart_(\d+)', svg_file.stem)
                if match:
                    structure["extracted_data"]["charts"].append({
                        "page_number": int(match.group(1)),
                        "svg_path": str(svg_file)
                    })

        return structure

    def _add_metadata_section(self, doc, structured_data: Dict):
        """Add metadata section to document."""
        from docx.shared import Pt

        metadata = structured_data.get("metadata", {})

        if metadata.get("author") or metadata.get("created_date"):
            p = doc.add_paragraph()
            p.add_run("Document Information").bold = True
            p.add_run("\n")

            if metadata.get("author"):
                p.add_run(f"Author: {metadata['author']}\n")
            if metadata.get("created_date"):
                p.add_run(f"Created: {metadata['created_date']}\n")
            if metadata.get("page_count"):
                p.add_run(f"Pages: {metadata['page_count']}\n")

            doc.add_paragraph()  # Spacer

    def _add_text_content(self, doc, text: str):
        """Add text content, preserving basic structure."""
        from docx.shared import Pt

        # Split into paragraphs
        paragraphs = text.split('\n\n')

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Check if it looks like a heading
            if self._is_heading(para_text):
                level = self._detect_heading_level(para_text)
                clean_text = para_text.lstrip('#').strip()
                doc.add_heading(clean_text, level=min(level, 4))
            else:
                p = doc.add_paragraph(para_text)

    def _is_heading(self, text: str) -> bool:
        """Check if text appears to be a heading."""
        # Markdown headings
        if text.startswith('#'):
            return True
        # ALL CAPS short text
        if text.isupper() and len(text) < 100:
            return True
        # Numbered headings
        if re.match(r'^\d+\.?\s+[A-Z]', text) and len(text) < 100:
            return True
        return False

    def _detect_heading_level(self, text: str) -> int:
        """Detect heading level."""
        if text.startswith('#'):
            level = len(text) - len(text.lstrip('#'))
            return min(level, 4)
        return 2

    def _add_table(self, doc, table_data: Dict):
        """Add a table to the document."""
        from docx.shared import Inches, Pt

        headers = table_data.get("headers", [])
        data = table_data.get("data", [])

        if not headers or not data:
            return

        # Add table caption
        p = doc.add_paragraph()
        p.add_run(f"Table (Page {table_data.get('page_number', '?')})").italic = True

        # Create table
        num_cols = len(headers)
        num_rows = len(data) + 1  # +1 for header

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = str(header)
            # Bold header text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for row_idx, row_data in enumerate(data):
            row = table.rows[row_idx + 1]
            for col_idx, header in enumerate(headers):
                cell_value = row_data.get(header, "")
                row.cells[col_idx].text = str(cell_value)

        doc.add_paragraph()  # Spacer

    def _add_chart(self, doc, chart_data: Dict, document_dir: Path):
        """Add a chart (SVG converted to image) to the document."""
        from docx.shared import Inches

        svg_path = chart_data.get("svg_path")
        if not svg_path:
            return

        svg_path = Path(svg_path)
        if not svg_path.is_absolute():
            svg_path = document_dir / svg_path

        if not svg_path.exists():
            return

        # Convert SVG to PNG for Word embedding
        png_data = self._svg_to_png(svg_path)
        if not png_data:
            return

        # Add caption
        p = doc.add_paragraph()
        p.add_run(f"Chart (Page {chart_data.get('page_number', '?')})").italic = True

        # Add image from bytes
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(png_data)
            tmp_path = tmp.name

        try:
            doc.add_picture(tmp_path, width=Inches(5))
        finally:
            Path(tmp_path).unlink(missing_ok=True)

        doc.add_paragraph()  # Spacer

    def _svg_to_png(self, svg_path: Path) -> Optional[bytes]:
        """Convert SVG to PNG bytes."""
        try:
            import cairosvg
            return cairosvg.svg2png(url=str(svg_path))
        except ImportError:
            logger.warning("cairosvg not installed, SVG charts will be skipped")
            return None
        except Exception as e:
            logger.warning(f"SVG to PNG conversion failed: {e}")
            return None

    def _add_regenerated_images(self, doc, page_num: int, document_dir: Path):
        """Add regenerated images for a page."""
        from docx.shared import Inches

        regen_dir = document_dir / "regenerated"
        if not regen_dir.exists():
            return

        pattern = f"regen_page_{page_num:03d}_*.png"
        for img_path in sorted(regen_dir.glob(pattern)):
            p = doc.add_paragraph()
            p.add_run(f"Image (Page {page_num}, Regenerated)").italic = True

            try:
                doc.add_picture(str(img_path), width=Inches(4))
            except Exception as e:
                logger.warning(f"Failed to add image {img_path}: {e}")

            doc.add_paragraph()  # Spacer

    def _add_image(self, doc, image_data: Dict, document_dir: Path):
        """Add an original extracted image."""
        from docx.shared import Inches

        file_path = image_data.get("file_path")
        if not file_path:
            return

        file_path = Path(file_path)
        if not file_path.is_absolute():
            file_path = document_dir / file_path

        if not file_path.exists():
            return

        p = doc.add_paragraph()
        p.add_run(f"Image (Page {image_data.get('page_number', '?')})").italic = True

        try:
            doc.add_picture(str(file_path), width=Inches(4))
        except Exception as e:
            logger.warning(f"Failed to add image {file_path}: {e}")

        doc.add_paragraph()  # Spacer

    def export_from_structured_document(
        self,
        structured_doc,
        output_path: str | Path,
        document_dir: Optional[str | Path] = None
    ) -> str:
        """
        Export from a StructuredDocument object.

        Args:
            structured_doc: StructuredDocument object
            output_path: Path for output .docx file
            document_dir: Optional directory for resolving relative paths

        Returns:
            Path to the created Word document
        """
        data = structured_doc.to_dict()

        if document_dir:
            return self.export(document_dir, output_path, structured_data=data)
        else:
            # Create temp dir structure
            import tempfile
            with tempfile.TemporaryDirectory() as tmp_dir:
                return self.export(tmp_dir, output_path, structured_data=data)
